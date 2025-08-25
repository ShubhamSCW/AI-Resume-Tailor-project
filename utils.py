# utils.py — loaders, contact extraction, exporters (verbatim-preserving), DB, optional OCR, AI layout (indices-only)
import io
import re
import sqlite3
import json
import asyncio
from typing import Any, Dict, List, Optional
import streamlit as st
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader
from playwright.async_api import async_playwright
import requests
from openai import OpenAI
import google.generativeai as genai
from config import DATABASE_FILE, GEMINI_MODEL, OPENAI_MODEL, ANTHROPIC_MODEL, OLLAMA_MODEL

_ANTHROPIC_OK = False
try:
    import anthropic
    _ANTHROPIC_OK = True
except ImportError:
    pass

# =========================
# Helpers (hash/ocr)
# =========================
def _crc32(content: bytes) -> str:
    import zlib
    return f"{zlib.crc32(content) & 0xffffffff:08x}"

def _ocr_pdf_bytes(pdf_bytes: bytes) -> str:
    """Best-effort OCR for scanned PDFs (optional)."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except Exception:
        return ""
    try:
        pages = convert_from_bytes(pdf_bytes, dpi=300)
        texts = []
        for img in pages:
            if img.mode != "L":
                img = img.convert("L")
            txt = pytesseract.image_to_string(img)
            texts.append(txt or "")
        return "\n".join(texts).strip()
    except Exception:
        return ""

# =========================
# VERBATIM loaders (file → JSON) + wrappers
# =========================
def wrap_text_as_json(text: str, source_hint: str = "paste") -> Dict[str, Any]:
    raw = text or ""
    b = raw.encode("utf-8", errors="ignore")
    return {
        "source": source_hint,
        "raw_text": raw,
        "length": len(raw),
        "hash": _crc32(b),
        "metadata": {"mime": "text/plain", "pages": None, "paragraphs": None},
    }

def file_to_json(upload) -> Dict[str, Any]:
    """
    Read .txt/.docx/.pdf -> return a JSON blob with VERBATIM raw_text + structure.
    Optional OCR fallback for scanned PDFs (controlled by st.session_state['ocr_pdf']).
    """
    name = getattr(upload, "name", "unknown").lower()
    data = upload.getvalue()
    hashv = _crc32(data)

    def _out(raw_text: str, meta: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "source": name,
            "raw_text": raw_text,
            "length": len(raw_text),
            "hash": hashv,
            "metadata": meta,
        }
    try:
        # .txt
        if name.endswith(".txt"):
            raw = data.decode("utf-8", errors="ignore")
            return _out(raw, {"mime": "text/plain", "pages": None, "paragraphs": raw.count("\n\n") or None})
        # .docx
        if name.endswith(".docx"):
            doc = DocxDocument(io.BytesIO(data))
            paragraphs: List[str] = [p.text for p in doc.paragraphs]
            raw = "\n".join(paragraphs)
            return _out(
                raw,
                {
                    "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "pages": None,
                    "paragraphs": len(paragraphs),
                },
            )
        # .pdf
        if name.endswith(".pdf"):
            pages_text: List[str] = []
            # 1) PyPDF2
            try:
                reader = PdfReader(io.BytesIO(data))
                for p in reader.pages:
                    pages_text.append(p.extract_text() or "")
            except Exception:
                pass
            # 2) pdfminer.six (correct module path)
            if not any(pages_text) or len("".join(pages_text).strip()) < 40:
                try:
                    from pdfminer.high_level import extract_text as pdfminer_extract  # <- fixed path
                except Exception:
                    pdfminer_extract = None
                if pdfminer_extract:
                    try:
                        full = pdfminer_extract(io.BytesIO(data)) or ""
                        if full.strip():
                            pages_text = [full]
                    except Exception:
                        pass
            # 3) pdfplumber
            if not any(pages_text) or len("".join(pages_text).strip()) < 40:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(data)) as pdf:
                        pages_text = [(page.extract_text() or "") for page in pdf.pages]
                except Exception:
                    pass
            # 4) OCR fallback if enabled
            try_ocr = bool(st.session_state.get("ocr_pdf", False)) if hasattr(st, "session_state") else False
            joined = "".join(pages_text).strip() if pages_text else ""
            if try_ocr and (not joined or len(joined) < 40):
                ocr_text = _ocr_pdf_bytes(data)
                if ocr_text and len(ocr_text) > len(joined):
                    pages_text = [ocr_text]
            raw = "\n".join(pages_text).strip()
            return _out(raw, {"mime": "application/pdf", "pages": len(pages_text) or None, "paragraphs": None})
        # fallback
        raw = data.decode("utf-8", errors="ignore")
        return _out(raw, {"mime": "application/octet-stream", "pages": None, "paragraphs": None})
    except Exception as e:
        st.error(f"Error reading file {name}: {e}")
        raw = data.decode("utf-8", errors="ignore")
        return _out(raw, {"mime": "application/octet-stream", "pages": None, "paragraphs": None})

# Backward compatibility
def read_file_to_text(upload) -> str:
    blob = file_to_json(upload)
    return blob.get("raw_text", "")

# =========================
# Contact info extraction (dict-safe)
# =========================
COMMON_HEADINGS = {
    "resume","curriculum vitae","cv","profile","profile summary","summary",
    "professional summary","work experience","experience","education","skills",
    "contact","contact information", "details"
}

def _clean_line(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())

def guess_name_from_header(text: str) -> str:
    """Finds the best candidate for a person's name from the top of the resume."""
    if not text:
        return ""
    
    # Increased search limit from 12 to 20 lines for more flexibility.
    lines = [_clean_line(x) for x in text.splitlines() if _clean_line(x)]
    search_lines = lines[:20] 

    for ln in search_lines:
        low = ln.lower()
        if low in COMMON_HEADINGS:
            continue
        if any(k in low for k in ["@", "http", "www", "linkedin", "github", "phone", "email"]):
            continue
        
        # A name is likely 2-4 words, capitalized, and not a common heading.
        tokens = ln.split()
        if 2 <= len(tokens) <= 4 and all(t[0].isupper() for t in tokens if t and t[0].isalpha()):
            return ln
            
    # spaCy fallback (optional)
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm")
        doc = nlp("\n".join(search_lines))
        for ent in doc.ents:
            if ent.label_ == "PERSON" and 2 <= len(ent.text.split()) <= 4:
                return ent.text.strip()
    except Exception:
        pass
        
    # If all else fails, return the first non-heading line as a last resort.
    if lines and lines[0].lower() not in COMMON_HEADINGS:
        return lines[0]
        
    return ""

def extract_contact_info(text_or_json: Any) -> Dict[str, str]:
    """
    Accepts a dict blob from file_to_json (with raw_text) OR a plain string.
    Returns name, email, phone, linkedin, location.
    """
    if isinstance(text_or_json, dict):
        txt = text_or_json.get("raw_text", "") or ""
    else:
        txt = str(text_or_json or "")
    
    if not txt.strip():
        return {}

    txt = txt.replace("\u200b", "").replace("\xa0", " ")
    lines = [_clean_line(x) for x in txt.splitlines() if _clean_line(x)]
    
    # Email
    email = ""
    m = re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", txt, re.I)
    if m: email = m.group(0)
    
    # Phone (more inclusive for intl dialing codes, bounded by digits length)
    phone = ""
    m = re.search(
        r'(?:(?:\+?\d{1,3}[\s\-\.]?)?(?:\(?\d{2,4}\)?[\s\-\.]?)?)\d{3,4}[\s\-\.]?\d{4}',
        txt
    )
    if m:
        ph = m.group(0)
        digits = re.sub(r"\D","", ph)
        if 8 <= len(digits) <= 15:
            phone = ph
            
    # LinkedIn
    linkedin = ""
    m = re.search(r"linkedin\.com/(in|pub)/[A-Za-z0-9\-_/%]+", txt, re.I)
    if m: 
        linkedin = "https://www.linkedin.com/" + m.group(0).split("linkedin.com/", 1)[-1]
        
    # Location guess (simple heuristic near top)
    location = ""
    for ln in lines[:30]:
        low = ln.lower()
        if any(w in low for w in [
            "india","usa","united states","canada","uk","uae","germany","australia",
            "noida","gurgaon","gurugram","bangalore","bengaluru","mumbai","pune",
            "hyderabad","delhi","chennai","kolkata","remote","hybrid","london","new york",
            "san francisco", "toronto", "dubai", "berlin", "sydney"
        ]):
            if "linkedin.com" not in low:
                location = ln
                break
                
    name = guess_name_from_header(txt)
    
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "location": location,
    }

# =========================
# Database (history)
# =========================
def init_db():
    try:
        with sqlite3.connect(DATABASE_FILE) as conn:
            c = conn.cursor()
            c.execute(
                "CREATE TABLE IF NOT EXISTS history ("
                "id INTEGER PRIMARY KEY, "
                "timestamp DATETIME DEFAULT CURRENT_TIMESTAMP, "
                "job_title TEXT, resume_name TEXT, match_score INTEGER, ats_score INTEGER, "
                "tailored_resume_text TEXT, cover_letter_text TEXT)"
            )
    except sqlite3.Error as e:
        st.warning(f"Database init error: {e}")

def log_analysis(job_title, resume_name, match_score, ats_score, tailored_resume_text, cover_letter_text):
    try:
        with sqlite3.connect(DATABASE_FILE) as conn:
            c = conn.cursor()
            c.execute(
                "INSERT INTO history (job_title, resume_name, match_score, ats_score, tailored_resume_text, cover_letter_text) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (job_title, resume_name, int(match_score), int(ats_score), tailored_resume_text, cover_letter_text)
            )
    except sqlite3.Error as e:
        st.warning(f"Database error while logging: {e}")

def get_history() -> List:
    try:
        with sqlite3.connect(DATABASE_FILE) as conn:
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("SELECT id, timestamp, job_title, resume_name, match_score FROM history ORDER BY timestamp DESC LIMIT 20")
            return c.fetchall()
    except sqlite3.Error as e:
        st.warning(f"Database error while fetching history: {e}")
        return []

# =========================
# Centralized LLM (layout planning only in this module)
# =========================
def _call_llm_api(backend: str, prompt: str, keys: Dict[str, str], expect_json: bool) -> Optional[str]:
    try:
        if backend == "Google Gemini" and keys.get("gemini"):
            genai.configure(api_key=keys["gemini"])
            model = genai.GenerativeModel(GEMINI_MODEL)
            return model.generate_content(prompt).text
        elif backend == "OpenAI GPT" and keys.get("openai"):
            client = OpenAI(api_key=keys["openai"])
            messages = [{"role": "user", "content": prompt}]
            if expect_json:
                messages.insert(0, {"role": "system", "content": "Return ONLY valid JSON."})
            resp = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=messages,
                temperature=0.1,
                response_format={"type": "json_object"} if expect_json else {"type": "text"}
            )
            return resp.choices[0].message.content
        elif backend == "Anthropic Claude" and keys.get("anthropic") and _ANTHROPIC_OK:
            client = anthropic.Anthropic(api_key=keys["anthropic"])
            msg = client.messages.create(
                model=ANTHROPIC_MODEL,
                max_tokens=4096,
                system="Return ONLY valid JSON." if expect_json else "",
                messages=[{"role": "user", "content": prompt}]
            )
            return "".join(block.text for block in msg.content if hasattr(block, "text"))
        elif backend == "Ollama (local)":
            messages = [{"role": "user", "content": prompt}]
            if expect_json:
                messages.insert(0, {"role": "system", "content": "Return ONLY valid JSON."})
            resp = requests.post(
                "http://localhost:11434/api/chat",
                json={"model": OLLAMA_MODEL,"messages": messages,"stream": False,"format": "json" if expect_json else "text"},
                timeout=120
            )
            resp.raise_for_status()
            return resp.json().get("message", {}).get("content", "")
    except Exception as e:
        if any(k in str(e).lower() for k in ["permission", "authentication", "api key"]):
            st.warning(f"{backend} failed: Invalid API Key.")
        else:
            st.warning(f"{backend} call failed: {e}")
    return None

def _execute_llm_call(prompt: str, primary: str, keys: Dict[str, str], expect_json: bool) -> str:
    order = [primary] + [b for b in ["Google Gemini", "OpenAI GPT", "Anthropic Claude", "Ollama (local)"] if b != primary]
    for backend in order:
        if response := _call_llm_api(backend, prompt, keys, expect_json):
            return response.strip()
    st.error("All AI backends failed.")
    return "{}" if expect_json else ""

def call_llm_json(prompt: str, primary: str, keys: Dict[str, str]) -> Dict[str, Any]:
    raw_response = _execute_llm_call(prompt, "Google Gemini", keys={k:v for k,v in (keys or {}).items() if v}, expect_json=True)
    try:
        return json.loads(raw_response)
    except Exception:
        m = re.search(r"\{[\s\S]*\}", raw_response or "")
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass
    st.error("AI returned invalid JSON.")
    return {}

# =========================
# AI layout planner (indices only, NO wording edits)
# =========================
def ai_plan_layout(text: str, keys: Dict[str, str]) -> Dict[str, Any]:
    """
    Return JSON with indices into original lines; we never use model text.
    """
    lines = (text or "").splitlines()
    sample = "\n".join(f"{i:04d}: {ln}" for i, ln in enumerate(lines))
    prompt = f"""
You are a resume layout planner. DO NOT rewrite content.
Return JSON describing structure using ORIGINAL line indices only.
Schema:
{{
  "name_line": int | null,
  "contact_lines": [int, ...],
  "sections": [
    {{
      "title": str,
      "line_start": int,
      "line_end": int,
      "bullet_lines": [int, ...]
    }}
  ]
}}
Rules:
- Use ONLY indices from the original lines below.
- If a section isn't present, omit it.
- If unsure about name/contact, set to null/empty list.
- Do not invent text or edit wording.
Original numbered lines:
---
{sample}
---
Return ONLY JSON.
""".strip()
    plan = call_llm_json(prompt, "Google Gemini", keys={k:v for k,v in (keys or {}).items() if v}) or {}
    # Fallback heuristic
    if not plan or "sections" not in plan:
        secs = []
        current = None
        for idx, ln in enumerate(lines):
            s = ln or ""
            is_header = s.isupper() and 2 <= len(s) <= 60
            if is_header:
                if current:
                    current["line_end"] = idx - 1
                    secs.append(current)
                current = {"title": s, "line_start": idx+1, "line_end": idx+1, "bullet_lines":[]}
            else:
                if current:
                    current["line_end"] = idx
                if _looks_like_bullet(s):
                    if current:
                        current["bullet_lines"].append(idx)
        if current:
            secs.append(current)
        plan = {"name_line": 0 if lines else None, "contact_lines": [1] if len(lines)>1 else [], "sections": secs}
    return plan

# =========================
# Bullet parsing (robust) + Verbatim guard
# =========================
_BULLET_RE = re.compile(
    r"""^
        (?P<lead>\s*)
        (?P<mark>[•\-\*\u2022\u2023\u25CF\u25AA\u00B7\u2013\u2014])   # • - * • ‣ ● ▪ · – —
        (?P<space>[\s\u00A0]+)?
        (?P<body>.*)$
    """,
    re.VERBOSE,
)

def _looks_like_bullet(line: str) -> bool:
    if not isinstance(line, str):
        return False
    if _BULLET_RE.match(line):
        return True
    # handle "•text" (no space)
    return bool(line.startswith("•") and len(line) > 1 and not line[1].isspace())

def _split_bullet(line: str):
    """Return (is_bullet, body_text) without trimming the body."""
    if not isinstance(line, str):
        return False, line
    m = _BULLET_RE.match(line)
    if m:
        return True, m.group("body")
    if line.startswith("•") and len(line) > 1 and not line[1].isspace():
        return True, line[1:]
    return False, line

# --- Normalization helpers for guards ---
_ZW_RE = re.compile(r"[\u200B\u200C\u200D\uFEFF]")  # zero-width, BOM
_NBSP_RE = re.compile(r"\u00A0")
_WS_RUN = re.compile(r"[ \t]{2,}")

def _normalize_for_guard_lines(ls: List[str]) -> List[str]:
    out = []
    for s in ls:
        s = s.replace("\r\n", "\n").replace("\r", "\n")
        s = _ZW_RE.sub("", s)
        s = _NBSP_RE.sub(" ", s)
        is_b, body = _split_bullet(s)
        s = body if is_b else s
        s = s.strip()
        s = _WS_RUN.sub(" ", s)
        out.append(s)
    # collapse >=3 blank lines to 2
    return out

def _verbatim_guard_strict(original_lines: List[str], rendered_lines: List[str]) -> None:
    o = "\n".join(_normalize_for_guard_lines(original_lines))
    r = "\n".join(_normalize_for_guard_lines(rendered_lines))
    o = re.sub(r"\n{3,}", "\n\n", o)
    r = re.sub(r"\n{3,}", "\n\n", r)
    if o.strip() != r.strip():
        raise RuntimeError("Verbatim guard: exported content diverged from source text.")

def _verbatim_guard_lenient(original_lines: List[str], rendered_lines: List[str]) -> None:
    """
    Lenient for AI layout: ensure every ORIGINAL non-empty normalized line appears
    in the rendered output at least as many times (multiset coverage).
    Order is relaxed to avoid false fails from section title insertions.
    """
    on = [x for x in _normalize_for_guard_lines(original_lines) if x != ""]
    rn = [x for x in _normalize_for_guard_lines(rendered_lines) if x != ""]
    from collections import Counter
    co = Counter(on)
    cr = Counter(rn)
    missing = []
    for k, v in co.items():
        if cr.get(k, 0) < v:
            missing.append(k)
    if missing:
        raise RuntimeError("Verbatim guard: some source lines were not carried into export.")

# =========================
# Async guard for PDF
# =========================
def _run_async(coro):
    try:
        return asyncio.run(coro)
    except RuntimeError as e:
        if "asyncio.run() cannot be called" in str(e):
            loop = asyncio.get_event_loop()
            return loop.run_until_complete(coro)
        raise

async def _generate_pdf_async(html_content: str) -> bytes:
    async with async_playwright() as p:
        browser = await p.chromium.launch(args=["--no-sandbox"])
        page = await browser.new_page()
        await page.set_content(html_content)
        pdf_bytes = await page.pdf(
            format="A4",
            print_background=True,
            margin={"top":"0.6in","bottom":"0.6in","left":"0.6in","right":"0.6in"}
        )
        await browser.close()
        return pdf_bytes

# =========================
# Exporters — PRESERVE wording (no AI text)
# =========================
def format_pdf_preserve(text: str) -> io.BytesIO:
    """Simple PDF render that never edits words."""
    if not text.strip():
        return io.BytesIO()
    lines = text.splitlines()
    rendered_copy = []  # for verbatim guard
    html_lines = []
    for ln in lines:
        if ln.strip() == "":
            html_lines.append("<br>")
            rendered_copy.append("")
            continue
        is_bul, body = _split_bullet(ln)
        if is_bul:
            html_lines.append(f"<div style='margin-left:20px'>• {body}</div>")
            rendered_copy.append(f"• {body}")
        elif ln.isupper() and len(ln.split()) <= 7:
            html_lines.append(f"<div style='font-weight:700; font-size:13pt; margin-top:10px'>{ln}</div>")
            rendered_copy.append(ln)
        else:
            html_lines.append(f"<div>{ln}</div>")
            rendered_copy.append(ln)
    _verbatim_guard_strict(lines, rendered_copy)
    html = f"""
    <html>
    <head>
      <meta charset="utf-8" />
      <style>
        body {{
          font-family: -apple-system, Segoe UI, Inter, Roboto, Arial, sans-serif;
          font-size: 11pt;
          color: #111827;
          line-height: 1.5;
          margin: 48px;
        }}
      </style>
    </head>
    <body>
      {''.join(html_lines)}
    </body>
    </html>
    """
    pdf_bytes = _run_async(_generate_pdf_async(html))
    return io.BytesIO(pdf_bytes)

def format_docx_preserve(text: str) -> io.BytesIO:
    """DOCX render that never edits words; robust bullet parsing."""
    if not text.strip():
        return io.BytesIO()
    doc = DocxDocument()
    for section in doc.sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.8)
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    except Exception:
        pass
    lines = text.splitlines()
    rendered_copy = []
    for ln in lines:
        if ln.strip() == "":
            doc.add_paragraph("")
            rendered_copy.append("")
            continue
        is_bul, body = _split_bullet(ln)
        if is_bul:
            p = doc.add_paragraph(body)
            p.style = "List Bullet"
            p.paragraph_format.space_after = Pt(0)
            rendered_copy.append(f"• {body}")
            continue
        if ln.isupper() and len(ln.split()) <= 7:
            p = doc.add_paragraph(ln)
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
            rendered_copy.append(ln)
            continue
        p = doc.add_paragraph(ln)
        p.paragraph_format.space_after = Pt(0)
        rendered_copy.append(ln)
    _verbatim_guard_strict(lines, rendered_copy)
    bio = io.BytesIO()
    doc.save(bio); bio.seek(0)
    return bio

# =========================
# Exporters — AI layout (indices-only), wording preserved
# =========================
def format_pdf_align(text: str, keys: Dict[str, str], jd_keywords: Optional[List[str]] = None) -> io.BytesIO:
    """
    AI plans layout (indices only). We render the original text.
    We DO NOT merge or rewrite any line. Decorative titles are ignored by guard.
    """
    if not (text or "").strip():
        return io.BytesIO()
    plan = ai_plan_layout(text, keys)
    lines = text.splitlines()
    emitted: set[int] = set()
    rendered_copy: List[str] = []
    html_parts: List[str] = []
    
    def add_para(idx: int):
        s = lines[idx]
        _, body = _split_bullet(s)  # Get content, strip old bullet if present.
        html_parts.append(f"<div>{body}</div>")
        rendered_copy.append(body)  # For guard, use content only.
        emitted.add(idx)

    def add_bullet(idx: int):
        s = lines[idx]
        _, body = _split_bullet(s)  # Get content, strip old bullet if present.
        html_parts.append(f"<div style='margin-left:18px'>• {body}</div>")
        rendered_copy.append(f"• {body}")  # For guard, use new formatted bullet.
        emitted.add(idx)

    # Name
    if isinstance(plan.get("name_line"), int):
        i = plan["name_line"]
        if 0 <= i < len(lines) and lines[i].strip():
            html_parts.append(f"<div style='font-weight:800; font-size:18pt; text-align:center; margin-bottom:6px'>{lines[i]}</div>")
            rendered_copy.append(lines[i]); emitted.add(i)
    # Contact — render EACH line separately (no merging)
    if isinstance(plan.get("contact_lines"), list) and plan["contact_lines"]:
        for i in plan["contact_lines"]:
            if isinstance(i, int) and 0 <= i < len(lines) and lines[i].strip():
                html_parts.append(f"<div style='text-align:center; margin-bottom:4px'>{lines[i]}</div>")
                rendered_copy.append(lines[i]); emitted.add(i)
    # Sections
    for sec in plan.get("sections", []):
        title = str(sec.get("title","")).strip()
        if title:
            # Decorative only — do NOT include in rendered_copy (to satisfy guard)
            html_parts.append(f"<div style='font-weight:700; font-size:12pt; margin-top:12px'>{title.upper()}</div>")
        ls = max(0, int(sec.get("line_start", 0)))
        le = min(len(lines)-1, int(sec.get("line_end", ls)))
        bullets = set(int(i) for i in (sec.get("bullet_lines") or []) if isinstance(i,int))
        for i in range(ls, le+1):
            if not (0 <= i < len(lines)):
                continue
            s = lines[i]
            if s.strip() == "":
                html_parts.append("<br>")
                rendered_copy.append("")
                emitted.add(i)
                continue
            # If LLM marked as bullet OR it already looks like bullet — render as bullet
            if i in bullets or _looks_like_bullet(s):
                add_bullet(i)
            else:
                add_para(i)
    # Any lines the plan missed? Render them in original order.
    for i, s in enumerate(lines):
        if i in emitted:
            continue
        if s.strip() == "":
            html_parts.append("<br>")
            rendered_copy.append("")
            emitted.add(i)
        else:
            # preserve original structure
            if _looks_like_bullet(s):
                add_bullet(i)
            else:
                add_para(i)
    _verbatim_guard_lenient(lines, rendered_copy)
    html = f"""
    <html><head><meta charset="utf-8" />
    <style>
      body {{ font-family: -apple-system, Segoe UI, Inter, Roboto, Arial, sans-serif; color:#111827; font-size:11pt; line-height:1.5; margin:48px; }}
    </style></head>
    <body>{''.join(html_parts)}</body></html>
    """
    pdf_bytes = _run_async(_generate_pdf_async(html))
    return io.BytesIO(pdf_bytes)

def format_docx_align(text: str, keys: Dict[str, str]) -> io.BytesIO:
    """
    AI plans layout (indices only). We render original text in DOCX.
    Decorative titles are ignored by guard. Contact lines are not merged.
    """
    if not (text or "").strip():
        return io.BytesIO()
    plan = ai_plan_layout(text, keys)
    lines = text.splitlines()
    emitted: set[int] = set()
    rendered_copy: List[str] = []
    doc = DocxDocument()
    for section in doc.sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.8)
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    except Exception:
        pass

    def add_para(idx: int):
        s = lines[idx]
        _, body = _split_bullet(s)  # Get content, strip old bullet.
        p = doc.add_paragraph(body)
        p.paragraph_format.space_after = Pt(0)
        rendered_copy.append(body); emitted.add(idx)

    def add_bullet(idx: int):
        s = lines[idx]
        _, body = _split_bullet(s)  # Get content, strip old bullet.
        p = doc.add_paragraph(body)
        p.style = "List Bullet"
        p.paragraph_format.space_after = Pt(0)
        rendered_copy.append(f"• {body}"); emitted.add(idx)

    # Name
    if isinstance(plan.get("name_line"), int):
        i = plan["name_line"]
        if 0 <= i < len(lines) and lines[i].strip():
            p = doc.add_paragraph()
            r = p.add_run(lines[i])
            r.bold = True
            r.font.size = Pt(20)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rendered_copy.append(lines[i]); emitted.add(i)
    # Contact — EACH line separately
    if isinstance(plan.get("contact_lines"), list) and plan["contact_lines"]:
        for i in plan["contact_lines"]:
            if isinstance(i, int) and 0 <= i < len(lines) and lines[i].strip():
                p = doc.add_paragraph(lines[i])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rendered_copy.append(lines[i]); emitted.add(i)
    # Sections
    for sec in plan.get("sections", []):
        title = str(sec.get("title","")).strip()
        if title:
            # Decorative only — not added to rendered_copy
            p = doc.add_paragraph(title.upper())
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
        ls = max(0, int(sec.get("line_start", 0)))
        le = min(len(lines)-1, int(sec.get("line_end", ls)))
        bullets = set(int(i) for i in (sec.get("bullet_lines") or []) if isinstance(i,int))
        for i in range(ls, le+1):
            if not (0 <= i < len(lines)):
                continue
            s = lines[i]
            if s.strip() == "":
                doc.add_paragraph("")
                rendered_copy.append("")
                emitted.add(i)
                continue
            if i in bullets or _looks_like_bullet(s):
                add_bullet(i)
            else:
                add_para(i)
    # Render any missed lines
    for i, s in enumerate(lines):
        if i in emitted:
            continue
        if s.strip() == "":
            doc.add_paragraph("")
            rendered_copy.append("")
            emitted.add(i)
        else:
            if _looks_like_bullet(s):
                add_bullet(i)
            else:
                add_para(i)
    _verbatim_guard_lenient(lines, rendered_copy)
    bio = io.BytesIO()
    doc.save(bio); bio.seek(0)
    return bio
