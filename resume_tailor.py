import streamlit as st
import google.generativeai as genai
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
import os
import tempfile

# -----------------------
# Resume Formatting Utils
# -----------------------
def format_docx(content, file_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Add Title
    doc.add_paragraph("Tailored Resume", "Title")

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Headings (like EXPERIENCE, SKILLS, etc.)
        if line.isupper() and len(line.split()) < 6:
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        # Bullet points
        if line.startswith(("-", "•", "*")):
            p = doc.add_paragraph(line[1:].strip(), style="List Bullet")
            run = p.runs[0]
            run.font.size = Pt(11)
            continue

        # Normal text
        p = doc.add_paragraph(line)
        run = p.runs[0]
        run.font.size = Pt(11)

    doc.save(file_path)


def format_pdf(content, file_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Headings
        if line.isupper() and len(line.split()) < 6:
            pdf.set_font("Arial", "B", 13)
            pdf.set_text_color(0, 51, 102)
            pdf.cell(0, 10, line, ln=True, align="L")
            pdf.set_font("Arial", size=11)
            pdf.set_text_color(0, 0, 0)
            continue

        # Bullet points
        if line.startswith(("-", "•", "*")):
            pdf.cell(10)
            pdf.multi_cell(0, 8, f"• {line[1:].strip()}")
            continue

        # Normal text
        pdf.multi_cell(0, 8, line)

    pdf.output(file_path)


# -----------------------
# AI Processing
# -----------------------
def tailor_resume(resume_text, jd_text, backend, api_key):
    prompt = f"""
    Rewrite and tailor the following resume to match the given Job Description.
    Keep it professional, ATS-friendly, cleanly formatted, and concise.
    Add relevant keywords from the JD naturally into the resume.

    Job Description:
    {jd_text}

    Resume:
    {resume_text}

    Provide only the improved resume content without markdown or extra symbols.
    """

    if backend == "Google Gemini":
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-pro")
        response = model.generate_content(prompt)
        return response.text.strip()

    elif backend == "OpenAI GPT":
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": "You are a professional resume writer."},
                      {"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()


# -----------------------
# Streamlit UI
# -----------------------
def main():
    st.set_page_config(page_title="AI Resume Tailor", layout="wide")

    st.title("📄 AI Resume Tailor")
    st.write("Upload your resume + JD and get a clean, tailored, ATS-friendly resume.")

    backend = st.radio("Select AI Backend:", ["Google Gemini", "OpenAI GPT"])
    api_key = st.text_input("Enter API Key", type="password")

    resume_file = st.file_uploader("Upload Resume (TXT or DOCX)", type=["txt", "docx"])
    jd_file = st.file_uploader("Upload Job Description (TXT or DOCX)", type=["txt", "docx"])

    if st.button("Generate Tailored Resume"):
        if not (resume_file and jd_file and api_key):
            st.error("Please upload both files and enter API key.")
        else:
            with st.spinner("Processing..."):
                import docx
                def read_file(file):
                    if file.name.endswith(".txt"):
                        return file.read().decode("utf-8")
                    elif file.name.endswith(".docx"):
                        doc = docx.Document(file)
                        return "\n".join([p.text for p in doc.paragraphs])

                resume_text = read_file(resume_file)
                jd_text = read_file(jd_file)

                tailored_text = tailor_resume(resume_text, jd_text, backend, api_key)

                # Save outputs
                tmpdir = tempfile.mkdtemp()
                docx_path = os.path.join(tmpdir, "Tailored_Resume.docx")
                pdf_path = os.path.join(tmpdir, "Tailored_Resume.pdf")

                format_docx(tailored_text, docx_path)
                format_pdf(tailored_text, pdf_path)

                st.success("✅ Resume tailored successfully!")

                with open(docx_path, "rb") as f:
                    st.download_button("⬇️ Download DOCX", f, file_name="Tailored_Resume.docx")

                with open(pdf_path, "rb") as f:
                    st.download_button("⬇️ Download PDF", f, file_name="Tailored_Resume.pdf")

                st.subheader("📌 Preview")
                st.text_area("Tailored Resume", tailored_text, height=400)


if __name__ == "__main__":
    try:
        main()
    except RuntimeError as e:
        if "ScriptRunContext" in str(e):
            print("⚠️ Streamlit warning: running in bare mode.")
            main()
        else:
            raise
