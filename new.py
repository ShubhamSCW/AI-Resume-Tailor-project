# app_ultimate.py
# -----------------------------------------------------------------------------------
# Ultimate AI Resume Suite (All-In-One) - v5.0 (ATS & Action Verb Analysis)
# Features:
# - Resume Parsing & Preprocessing
#   * Smarter python-docx styling (bold keywords, consistent spacing)
#   * NLP via spaCy (fallback to NLTK) + language detection
#   * Auto section detection (regex + optional AI assist)
#   * Basic multilingual parsing support (langdetect)
# - AI Enhancements
#   * Multiple backends: Gemini, OpenAI, Anthropic, Ollama (local), HuggingFace (optional)
#   * Automatic fallback chain if a backend fails
#   * Strict JSON outputs with granular, section-by-section feedback
#   * Readability (textstat) & complexity flag
# - Advanced Analytics <<< NEW >>>
#   * ATS (Applicant Tracking System) friendliness score and detailed checklist
#   * Action Verb and Quantitative Metrics counter to encourage results-driven language
# - UI/UX
#   * Drag & drop multi-upload
#   * Multi-resume comparison with improved interactive bar charts
#   * "Quick Metrics" dashboard for the top resume
#   * Interactive charts: keyword coverage, overall score, word clouds
#   * Session caching + optional SQLite persistence
#   * Light/Dark theme toggle
#   * Dedicated Cover Letter Tab with personalization
# - Exporting & Customization
#   * On-demand in-memory exports with robust heading detection
#   * Templates/styles (Minimalist, Modern, Creative)
#   * Dynamic highlight of missing keywords in tailored resume
# -----------------------------------------------------------------------------------

import os
import re
import io
import json
import time
import base64
import sqlite3
import tempfile
from typing import List, Dict, Any, Optional
from collections import Counter, defaultdict

import streamlit as st

# File parsing / formatting
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader

# PDF export (ReportLab)
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor

# NLP
import textstat
from langdetect import detect, LangDetectException

# Prefer spaCy; fallback to NLTK
_SPACY_OK = False
try:
    import spacy
    _SPACY_OK = True
except Exception:
    _SPACY_OK = False

_NLTK_OK = False
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    _NLTK_OK = True
except Exception:
    _NLTK_OK = False

# ML & viz
import plotly.graph_objects as go
import plotly.express as px
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Wordcloud
from wordcloud import WordCloud
import matplotlib.pyplot as plt

# Backends (optional installs)
import google.generativeai as genai
from openai import OpenAI
try:
    import anthropic
    _ANTHROPIC_OK = True
except Exception:
    _ANTHROPIC_OK = False

import requests

try:
    from transformers import pipeline
    _HF_OK = True
except Exception:
    _HF_OK = False


# ==============================================
# Helpers: text, parsing, NLP
# ==============================================

SECTION_PATTERNS = [
    r'^\s*(summary|professional summary|profile)\s*$',
    r'^\s*(experience|work experience|professional experience)\s*$',
    r'^\s*(education|academic background)\s*$',
    r'^\s*(skills|technical skills|core competencies)\s*$',
    r'^\s*(projects)\s*$',
    r'^\s*(certifications)\s*$',
    r'^\s*(publications)\s*$',
    r'^\s*(awards|achievements)\s*$',
    r'^\s*(volunteer|volunteering)\s*$',
]
SECTION_REGEX = re.compile("|".join(SECTION_PATTERNS), re.IGNORECASE | re.MULTILINE)

BUSINESS_JARGON_STOP_WORDS = {
    'experience', 'work', 'responsibilities', 'skills', 'company', 'team', 'project', 'role',
    'development', 'management', 'solution', 'solutions', 'technology', 'technologies',
    'environment', 'system', 'systems', 'tools', 'process', 'processes', 'data', 'analysis',
    'business', 'requirements', 'design', 'implementation', 'support', '-','–',
    'communication', 'years', 'job', 'description', 'candidate', 'ability', 'knowledge',
    'developer', 'engineer', 'specialist', 'analyst', 'architect', 'lead', 'manager',
    'devops', 'full', 'stack', 'software', 'data', 'scientist', 'product'
}

# <<< NEW FEATURE >>> List of strong action verbs for analysis
ACTION_VERBS = {
    'achieved', 'accelerated', 'administered', 'advised', 'advocated', 'analyzed', 'authored',
    'automated', 'built', 'calculated', 'centralized', 'chaired', 'coached', 'collaborated',
    'conceived', 'consolidated', 'constructed', 'consulted', 'converted', 'coordinated',
    'created', 'debugged', 'decreased', 'defined', 'delivered', 'designed', 'developed',
    'directed', 'documented', 'drove', 'eliminated', 'engineered', 'enhanced', 'established',
    'evaluated', 'executed', 'expanded', 'facilitated', 'founded', 'generated', 'grew',
    'guided', 'identified', 'implemented', 'improved', 'increased', 'influenced', 'initiated',
    'innovated', 'inspired', 'integrated', 'interpreted', 'invented', 'launched', 'led',
    'managed', 'mastered', 'mentored', 'modernized', 'motivated', 'negotiated', 'optimized',
    'orchestrated', 'overhauled', 'owned', 'pioneered', 'planned', 'prioritized', 'produced',
    'proposed', 'quantified', 'ran', 'rebuilt', 'reduced', 're-engineered', 'resolved',
    'restructured', 'revamped', 'saved', 'scaled', 'shipped', 'simplified', 'solved',
    'spearheaded', 'standardized', 'streamlined', 'strengthened', 'succeeded', 'supervised',
    'taught', 'trained', 'transformed', 'unified', 'won', 'wrote'
}

def detect_language(text: str) -> str:
    try:
        return detect(text)
    except LangDetectException:
        return "unknown"

def normalize_whitespace(text: str) -> str:
    text = re.sub(r'\r\n?', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

@st.cache_resource
def load_spacy_model(lang_code: str = "en"):
    if not _SPACY_OK:
        return None
    lang_map = {"en": "en_core_web_sm"}
    model_name = lang_map.get(lang_code, "en_core_web_sm")
    try:
        return spacy.load(model_name)
    except Exception:
        st.warning(f"spaCy model '{model_name}' not found. Please run: python -m spacy download {model_name}")
        return None

def extract_keywords_smarter(jd_text: str, top_k: int = 25) -> List[str]:
    if not jd_text:
        return []

    keywords = set()
    nlp = load_spacy_model("en")

    if nlp is None or not _SPACY_OK:
        st.warning("spaCy not available, falling back to basic keyword extraction.")
        tokens = re.findall(r"[A-Za-z][A-Za-z0-9+.#-]{2,}", jd_text.lower())
        return [word for word in tokens if word not in BUSINESS_JARGON_STOP_WORDS][:top_k]

    doc = nlp(jd_text)
    for token in doc:
        if token.pos_ in ['NOUN', 'PROPN']:
            keyword = token.text.lower().strip()
            if len(keyword) > 1 and keyword not in BUSINESS_JARGON_STOP_WORDS:
                keywords.add(keyword)

    text_lower = jd_text.lower()
    keyword_counts = {kw: text_lower.count(kw) for kw in keywords}

    sorted_keywords = sorted(keyword_counts.items(), key=lambda item: item[1], reverse=True)

    return [kw for kw, count in sorted_keywords[:top_k]]

def read_file_to_text(upload) -> str:
    name = upload.name.lower()
    try:
        if name.endswith('.txt'):
            return upload.read().decode('utf-8', errors='ignore')
        elif name.endswith('.docx'):
            doc = DocxDocument(upload)
            return "\n".join(p.text for p in doc.paragraphs)
        elif name.endswith('.pdf'):
            text = [p.extract_text() or "" for p in PdfReader(upload).pages]
            return "\n".join(text)
        return ""
    except Exception:
        return ""

def auto_detect_sections(text: str) -> Dict[str, List[str]]:
    blocks = defaultdict(list)
    current = "Other"
    for raw in text.split("\n"):
        line = raw.strip()
        if not line: continue
        if re.match(SECTION_REGEX, line):
            current = re.sub(r'\s+', ' ', line.title())
            blocks[current]
        else:
            blocks[current].append(line)
    return dict(blocks)

def compute_similarity(text_a: str, text_b: str) -> float:
    try:
        vect = TfidfVectorizer(stop_words='english')
        tfidf = vect.fit_transform([text_a, text_b])
        return float(cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0])
    except ValueError:
        return 0.0

def keyword_coverage(resume_text: str, jd_keywords: List[str]) -> Dict[str, int]:
    res_lc = resume_text.lower()
    return {k: int(k in res_lc) for k in jd_keywords}

def generate_wordcloud(words: List[str], title: str = ""):
    text = " ".join(words) if words else "no_data"
    wc = WordCloud(width=800, height=400, background_color="rgba(255, 255, 255, 0)", mode="RGBA", colormap='viridis').generate(text)
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=16)
    st.pyplot(fig, use_container_width=True)

# <<< NEW FEATURE >>> Function to analyze ATS friendliness
def analyze_ats_friendliness(text: str) -> Dict[str, Any]:
    checks = {}
    score = 100
    
    # Check 1: Contact Info
    has_email = bool(re.search(r'[\w\.-]+@[\w\.-]+', text))
    has_phone = bool(re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text))
    if has_email and has_phone:
        checks["✅ Contact Info"] = "Email and phone number found."
    else:
        checks["❌ Contact Info"] = "Missing email or phone number in a standard format."
        score -= 20

    # Check 2: Standard Sections
    found_sections = [s for s in ['experience', 'education', 'skills'] if re.search(f'^{s}', text, re.IGNORECASE | re.MULTILINE)]
    if len(found_sections) >= 3:
        checks["✅ Standard Sections"] = "Found essential sections (Experience, Education, Skills)."
    else:
        checks["❌ Standard Sections"] = "Missing one or more standard sections like 'Experience', 'Education', or 'Skills'."
        score -= 25

    # Check 3: Word Count
    word_count = len(text.split())
    if 400 <= word_count <= 800:
        checks["✅ Word Count"] = f"Good word count ({word_count}). Fits on 1-2 pages."
    else:
        checks["⚠️ Word Count"] = f"Word count is {word_count}. Aim for 400-800 words for most roles."
        score -= 10
    
    # Check 4: Parsable Format (heuristic)
    non_alpha_ratio = len(re.findall(r'[^a-zA-Z0-9\s]', text)) / len(text) if len(text) > 0 else 0
    if non_alpha_ratio < 0.05:
        checks["✅ Simple Formatting"] = "Appears to use standard characters, good for parsing."
    else:
        checks["⚠️ Simple Formatting"] = "High ratio of special characters detected. Avoid tables, columns, or graphics."
        score -= 15

    return {"score": max(0, score), "checks": checks}

# <<< NEW FEATURE >>> Function to count action verbs and metrics
def analyze_action_verbs_and_metrics(text: str) -> Dict[str, int]:
    words = set(re.findall(r'\b\w+\b', text.lower()))
    action_verb_count = len([v for v in words if v in ACTION_VERBS])
    
    # Regex to find numbers, percentages, dollar amounts
    metric_count = len(re.findall(r'\b\d+(\.\d+)?%?\b|\$\d+', text))
    
    return {"action_verbs": action_verb_count, "quantitative_metrics": metric_count}

def safe_json_extract(text: str) -> Dict[str, Any]:
    if not text:
        return {}
    try:
        return json.loads(text)
    except Exception:
        pass
    m = re.search(r"\{[\s\S]*\}", text)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return {}
    return {}

def call_gemini_json(prompt: str, api_key: str) -> Dict[str, Any]:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")
    resp = model.generate_content(prompt)
    return safe_json_extract(getattr(resp, "text", "") or "")

def call_openai_json(prompt: str, api_key: str) -> Dict[str, Any]:
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": "Return ONLY valid JSON."},
                  {"role": "user", "content": prompt}],
        temperature=0.2
    )
    return safe_json_extract(resp.choices[0].message.content)

def call_anthropic_json(prompt: str, api_key: str) -> Dict[str, Any]:
    if not _ANTHROPIC_OK:
        raise RuntimeError("Anthropic SDK not installed.")
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
        system="Return ONLY valid JSON."
    )
    text = "".join([blk.text for blk in msg.content if blk.type == "text"])
    return safe_json_extract(text)

def call_ollama_json(prompt: str) -> Dict[str, Any]:
    url = "http://localhost:11434/api/chat"
    data = {"model": "llama3.1", "messages": [{"role": "system", "content": "Return ONLY valid JSON."}, {"role": "user", "content": prompt}], "stream": False}
    r = requests.post(url, json=data, timeout=120)
    r.raise_for_status()
    content = r.json().get("message", {}).get("content", "")
    return safe_json_extract(content)

def call_huggingface_json(prompt: str) -> Dict[str, Any]:
    if not _HF_OK: raise RuntimeError("transformers not installed.")
    gen = pipeline("text-generation", model="gpt2")
    text = gen(prompt, max_length=512, num_return_sequences=1)[0]["generated_text"]
    return safe_json_extract(text)

def call_llm_json(prompt: str, primary: str, keys: Dict[str, str]) -> Dict[str, Any]:
    backends = ["Google Gemini", "OpenAI GPT", "Anthropic Claude", "Ollama (local)", "HuggingFace (local)"]
    order = [primary] + [b for b in backends if b != primary]
    last_err = None
    for b in order:
        try:
            if b == "Google Gemini" and keys.get("gemini"): return call_gemini_json(prompt, keys["gemini"])
            if b == "OpenAI GPT" and keys.get("openai"): return call_openai_json(prompt, keys["openai"])
            if b == "Anthropic Claude" and keys.get("anthropic"): return call_anthropic_json(prompt, keys["anthropic"])
            if b == "Ollama (local)": return call_ollama_json(prompt)
            if b == "HuggingFace (local)": return call_huggingface_json(prompt)
        except Exception as e:
            last_err = e
            st.warning(f"{b} failed: {e}")
            continue
    st.error(f"All backends failed. Last error: {last_err}")
    return {}

def call_llm_text(prompt: str, primary: str, keys: Dict[str, str]) -> str:
    order = [primary] + [b for b in ["Google Gemini", "OpenAI GPT", "Anthropic Claude", "Ollama (local)", "HuggingFace (local)"] if b != primary]
    for b in order:
        try:
            if b == "Google Gemini" and keys.get("gemini"):
                genai.configure(api_key=keys["gemini"])
                model = genai.GenerativeModel("gemini-1.5-flash")
                return (model.generate_content(prompt).text or "").strip()
            if b == "OpenAI GPT" and keys.get("openai"):
                client = OpenAI(api_key=keys["openai"])
                r = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.2)
                return r.choices[0].message.content.strip()
            if b == "Anthropic Claude" and keys.get("anthropic") and _ANTHROPIC_OK:
                client = anthropic.Anthropic(api_key=keys["anthropic"])
                msg = client.messages.create(model="claude-3-haiku-20240307", max_tokens=1500, messages=[{"role":"user","content":prompt}])
                return "".join([blk.text for blk in msg.content if getattr(blk, "type", "") == "text"]).strip()
            if b == "Ollama (local)":
                url = "http://localhost:11434/api/chat"
                data = {"model": "llama3.1", "messages":[{"role":"user","content":prompt}], "stream": False}
                r = requests.post(url, json=data, timeout=120)
                r.raise_for_status()
                return r.json().get("message", {}).get("content", "").strip()
            if b == "HuggingFace (local)" and _HF_OK:
                gen = pipeline("text-generation", model="gpt2")
                return gen(prompt, max_length=512, num_return_sequences=1)[0]["generated_text"].strip()
        except Exception as e:
            st.warning(f"{b} failed (text): {e}")
            continue
    return ""

# ==============================================
# Formatting & Export
# ==============================================

def style_docx_paragraph(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color

def format_docx_content(content: str, template: str = "Minimalist", keywords_to_bold: Optional[List[str]] = None):
    try:
        doc = DocxDocument()
        normal = doc.styles['Normal']
        normal.font.name = "Calibri" if template == "Minimalist" else ("Helvetica" if template == "Modern" else "Georgia")
        normal.font.size = Pt(11)
        accent = RGBColor(0x00, 0x33, 0x66) if template in ("Modern", "Creative") else None

        kws = [k.lower() for k in (keywords_to_bold or [])]

        for raw in content.split("\n"):
            line = raw.strip()
            if not line: continue

            if re.match(SECTION_REGEX, line) or (line.isupper() and len(line.split()) < 6):
                p = doc.add_paragraph()
                style_docx_paragraph(p.add_run(line), bold=True, size=13, color=accent)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith(("-", "•", "*")):
                p = doc.add_paragraph(style="List Bullet")
                text = line[1:].strip()
                tokens = re.split(r'(\W+)', text)
                for tok in tokens:
                    if kws and tok.lower() in kws:
                        style_docx_paragraph(p.add_run(tok), bold=True)
                    else:
                        p.add_run(tok)
            else:
                p = doc.add_paragraph()
                tokens = re.split(r'(\W+)', line)
                for tok in tokens:
                    if kws and tok.lower() in kws:
                        style_docx_paragraph(p.add_run(tok), bold=True)
                    else:
                        p.add_run(tok)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"Error creating DOCX file: {e}")
        return None

def format_pdf_content(content: str):
    try:
        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
        story = []
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Header', fontName='Helvetica-Bold', fontSize=12, spaceBefore=12, spaceAfter=6, textColor=HexColor('#003366')))
        styles.add(ParagraphStyle(name='CustomBullet', parent=styles['Normal'], leftIndent=20, firstLineIndent=0, spaceBefore=2))

        for raw in content.split("\n"):
            line = raw.strip()
            if not line: continue

            if re.match(SECTION_REGEX, line) or (line.isupper() and len(line.split()) < 6):
                p = Paragraph(line, styles['Header'])
            elif line.startswith(("-", "•", "*")):
                p = Paragraph(f"<bullet>&bull;</bullet>{line[1:].strip()}", styles['CustomBullet'])
            else:
                p = Paragraph(line, styles['Normal'])
            story.append(p)
        doc.build(story)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"Error creating PDF file: {e}")
        return None

def highlight_missing_keywords_html(text: str, missing: List[str]) -> str:
    if not missing: return f"<div>{text}</div>"
    repl = lambda m: f'<span style="background-color:#ffe4e1;border-radius:3px;padding:1px 2px">{m.group(0)}</span>' if m.group(0).lower() in missing else m.group(0)
    return re.sub(r"[A-Za-z][A-Za-z0-9+.#-]{1,}", repl, text, flags=re.IGNORECASE)

@st.cache_resource
def get_db(path="resume_suite.db"):
    conn = sqlite3.connect(path, check_same_thread=False)
    conn.execute("CREATE TABLE IF NOT EXISTS analyses(id INTEGER PRIMARY KEY, ts, backend, score, grade, payload TEXT)")
    return conn

def persist_analysis(conn, backend: str, score: int, grade: str, payload: Dict[str, Any]):
    conn.execute("INSERT INTO analyses(ts, backend, score, grade, payload) VALUES(?,?,?,?,?)",
                 (int(time.time()), backend, score, grade, json.dumps(payload)))
    conn.commit()

# ==============================================
# Streamlit App
# ==============================================

st.set_page_config(page_title="Ultimate AI Resume Suite", page_icon="🚀", layout="wide")

if 'theme' not in st.session_state:
    st.session_state.theme = False

is_dark = st.sidebar.toggle("🌙 Dark mode", value=st.session_state.theme, key="theme_toggle")
if is_dark:
    st.markdown("""<style> .stApp { background: #0e1117; color: #e0e0e0; } </style>""", unsafe_allow_html=True)

st.title("🚀 Ultimate AI Resume Suite")
st.caption("Analyze, compare, tailor, and export ATS-friendly resumes with multi-model AI and rich visuals.")

with st.sidebar.expander("⚙️ AI Configuration", expanded=True):
    primary_backend = st.selectbox("Primary Backend", ["Google Gemini", "OpenAI GPT", "Anthropic Claude", "Ollama (local)", "HuggingFace (local)"])
    gemini_key = st.text_input("Gemini API Key", type="password")
    openai_key = st.text_input("OpenAI API Key", type="password")
    anthropic_key = st.text_input("Anthropic API Key", type="password")
keys = {"gemini": gemini_key, "openai": openai_key, "anthropic": anthropic_key}

with st.sidebar.expander("🎨 Output Options"):
    candidate_name = st.text_input("Your Full Name", placeholder="e.g., Jane Doe") # <<< NEW FEATURE >>>
    resume_tone = st.selectbox("Resume Tone", ["Professional", "Creative", "Technical", "Enthusiastic"])
    docx_template = st.selectbox("DOCX Template", ["Minimalist", "Modern", "Creative"])
    persist_to_sqlite = st.checkbox("Persist analyses to SQLite")

st.subheader("1) Upload Resumes & Paste the Job Description")
c1, c2 = st.columns([1, 1])
with c1:
    resumes = st.file_uploader("Drag & drop resumes here", type=["txt", "docx", "pdf"], accept_multiple_files=True)
with c2:
    jd_text = st.text_area("Paste Job Description", height=220, placeholder="Paste JD here...")

if st.button("✨ Analyze & Generate", type="primary", use_container_width=True):
    if not resumes or not jd_text:
        st.warning("Please upload at least one resume and paste the JD.")
        st.stop()

    with st.spinner("Performing deep analysis and generating documents..."):
        jd_keywords = extract_keywords_smarter(jd_text, top_k=25)

        all_results = []
        for rf in resumes:
            text = read_file_to_text(rf)
            if not text:
                st.warning(f"Could not read content from {rf.name}. Skipping.")
                continue
            text_norm = normalize_whitespace(text)
            sim = compute_similarity(text_norm, jd_text)
            coverage = keyword_coverage(text_norm, jd_keywords)
            ats_analysis = analyze_ats_friendliness(text_norm) # <<< NEW FEATURE >>>
            verb_metric_analysis = analyze_action_verbs_and_metrics(text_norm) # <<< NEW FEATURE >>>

            # <<< MODIFIED >>> Improved AI prompt for more granular feedback
            grading_prompt = f"""
            Analyze the following resume against the job description. Return STRICT JSON with the following keys:
            - "SCORE": An integer score from 0 to 100 representing the overall match.
            - "GRADE": A letter grade from "A+" to "F".
            - "SUMMARY": A one-sentence summary of the candidate's fit.
            - "SECTION_FEEDBACK": A dictionary where keys are resume section titles (e.g., "Summary", "Experience", "Skills") and values are 1-2 bullet points of feedback for that section.

            Job Description (first 1000 chars): "{jd_text[:1000]}"
            Resume (first 1000 chars): "{text_norm[:1000]}"
            """
            result = call_llm_json(grading_prompt, primary_backend, keys) or {}
            score = int(result.get("SCORE", round(sim * 100)))
            grade = result.get("GRADE", "?")

            all_results.append({
                "name": rf.name, "text": text_norm, "score": score, "grade": grade,
                "summary": result.get("SUMMARY", ""),
                "section_feedback": result.get("SECTION_FEEDBACK", {}),
                "coverage": coverage,
                "ats": ats_analysis,
                "verbs_metrics": verb_metric_analysis
            })

        if not all_results:
            st.error("Analysis failed for all resumes.")
            st.stop()

        all_results.sort(key=lambda x: x["score"], reverse=True)
        st.session_state['analysis_results'] = all_results
        st.session_state['jd_keywords'] = jd_keywords
        st.session_state['jd_text'] = jd_text
        
        best_resume = all_results[0]
        tail_prompt = f"You are a professional resume writer. Return ONLY plain text. Rewrite and tailor this resume: '{best_resume['text']}' to perfectly match the job description: '{jd_text}' using a '{resume_tone}' tone. Ensure ATS-friendly formatting, action-driven bullets, and write ALL SECTION HEADINGS (like 'Experience', 'Skills') IN ALL CAPS."
        st.session_state['tailored_resume'] = call_llm_text(tail_prompt, primary_backend, keys)
        
        default_name = candidate_name or "the applicant"
        cover_prompt = f"Return ONLY plain text. Write a concise 3-4 paragraph cover letter for '{default_name}', based on this tailored resume: '{st.session_state['tailored_resume']}' and this job description: '{jd_text}'. Highlight 2-3 matching qualifications and end with a call to action."
        st.session_state['cover_letter'] = call_llm_text(cover_prompt, primary_backend, keys)

if 'analysis_results' in st.session_state:
    all_results = st.session_state['analysis_results']
    jd_keywords = st.session_state['jd_keywords']
    jd_text = st.session_state['jd_text']
    best_resume = all_results[0]

    st.success("Analysis and document generation complete!")
    
    tab_insights, tab_resume, tab_cover_letter, tab_export = st.tabs([
        "📊 Analysis Insights", 
        "✍️ View/Edit Tailored Resume", 
        "📄 View/Edit Cover Letter", 
        "💾 Export Documents"
    ])

    with tab_insights:
        # <<< NEW FEATURE >>> Quick Metrics Dashboard
        st.subheader(f"🏆 Quick Metrics for Top Resume: {best_resume['name']}")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Overall Score", f"{best_resume['score']}/100", delta=best_resume['grade'])
        m2.metric("ATS Score", f"{best_resume['ats']['score']}/100")
        m3.metric("Action Verbs", best_resume['verbs_metrics']['action_verbs'])
        m4.metric("Quant. Metrics", best_resume['verbs_metrics']['quantitative_metrics'])
        st.divider()

        st.subheader("Overall Match Score Comparison")
        score_df = pd.DataFrame(all_results).sort_values("score", ascending=True)
        fig_score = px.bar(score_df, x='score', y='name', orientation='h', title='Overall Match Score (Higher is Better)', labels={'score': 'Match Score (0-100)', 'name': 'Resume'}, text='score')
        fig_score.update_traces(texttemplate='%{text:.0f}', textposition='outside')
        fig_score.update_layout(yaxis={'categoryorder':'total ascending'})
        st.plotly_chart(fig_score, use_container_width=True)

        st.subheader("Keyword Coverage Score")
        for r in all_results:
            total_kws = len(jd_keywords)
            found_kws = sum(r['coverage'].values())
            r['coverage_percent'] = (found_kws / total_kws) * 100 if total_kws > 0 else 0
        coverage_df = pd.DataFrame(all_results).sort_values("coverage_percent", ascending=True)
        fig_coverage = px.bar(coverage_df, x='coverage_percent', y='name', orientation='h', title='Job Description Keyword Coverage', labels={'coverage_percent': 'Coverage (%)', 'name': 'Resume'}, text='coverage_percent')
        fig_coverage.update_traces(texttemplate='%{text:.0f}%', textposition='outside')
        fig_coverage.update_layout(xaxis_range=[0,105], yaxis={'categoryorder':'total ascending'})
        st.plotly_chart(fig_coverage, use_container_width=True)

        st.subheader("Per-Resume Details")
        for r in all_results:
            with st.expander(f"**{r['name']}** — Score: **{r['score']}** / Grade: **{r['grade']}**"):
                st.markdown(f"**AI Summary:** {r['summary']}")
                
                # <<< MODIFIED >>> Display granular section feedback
                if r['section_feedback']:
                    st.markdown("**AI Section-by-Section Feedback:**")
                    for section, feedback in r['section_feedback'].items():
                        st.markdown(f"- **{section}:** {' '.join(feedback) if isinstance(feedback, list) else feedback}")
                
                # <<< NEW FEATURE >>> Display detailed ATS and Verb/Metric analysis
                st.markdown(f"---")
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**ATS Friendliness: {r['ats']['score']}/100**")
                    for check, message in r['ats']['checks'].items():
                        st.markdown(f"<small>{check}: {message}</small>", unsafe_allow_html=True)
                with col2:
                    st.markdown("**Content Analysis:**")
                    st.markdown(f"- **Action Verbs Found:** {r['verbs_metrics']['action_verbs']}")
                    st.markdown(f"- **Quantitative Metrics Found:** {r['verbs_metrics']['quantitative_metrics']}")
                    st.markdown(f"- **Readability (Flesch):** {textstat.flesch_reading_ease(r['text']):.1f}")
                

    with tab_resume:
        st.subheader("View & Edit Your AI-Tailored Resume")
        st.info("The AI has rewritten your best resume to better match the job description. You can edit the text here before exporting.")
        
        tailored_resume_text = st.text_area(
            "AI-Tailored Resume", 
            value=st.session_state.get('tailored_resume', ""), 
            height=500, 
            key="tailed_text_area"
        )
        
        if tailored_resume_text:
            missing = sorted([k for k in jd_keywords if k not in tailored_resume_text.lower()])
            st.markdown("**Missing JD Keywords:**")
            st.info(", ".join(missing) if missing else "None Found! 🎉")
    
    with tab_cover_letter:
        st.subheader("View & Edit Your AI-Generated Cover Letter")
        st.info("The AI has written a cover letter based on your tailored resume. You can edit it here before exporting.")
        
        st.text_area(
            "AI-Generated Cover Letter", 
            value=st.session_state.get('cover_letter', ""), 
            height=500, 
            key="cover_letter_text_area"
        )

    with tab_export:
        st.subheader("Export Center")
        st.info("Download your generated documents in various formats.")
        
        tailored_text_export = st.session_state.get("tailed_text_area", "")
        cover_text_export = st.session_state.get("cover_letter_text_area", "")
        
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### Tailored Resume")
            if tailored_text_export:
                try:
                    docx_data = format_docx_content(tailored_text_export, docx_template, jd_keywords)
                    if docx_data:
                        st.download_button("⬇️ DOCX", data=docx_data, file_name="Tailored_Resume.docx", key="docx_res")
                except Exception as e:
                    st.error(f"DOCX Error: {e}")

                try:
                    pdf_data = format_pdf_content(tailored_text_export)
                    if pdf_data:
                        st.download_button("⬇️ PDF", data=pdf_data, file_name="Tailored_Resume.pdf", key="pdf_res")
                except Exception as e:
                    st.error(f"PDF Error: {e}")

                st.download_button("⬇️ TXT", data=tailored_text_export, file_name="Tailored_Resume.txt", key="txt_res")
            else:
                st.warning("No tailored resume available to export.")
        with c2:
            st.markdown("#### Cover Letter")
            if cover_text_export:
                try:
                    docx_data_cl = format_docx_content(cover_text_export, docx_template)
                    if docx_data_cl:
                        st.download_button("⬇️ DOCX", data=docx_data_cl, file_name="Cover_Letter.docx", key="docx_cl")
                except Exception as e:
                    st.error(f"DOCX Error: {e}")

                try:
                    pdf_data_cl = format_pdf_content(cover_text_export)
                    if pdf_data_cl:
                        st.download_button("⬇️ PDF", data=pdf_data_cl, file_name="Cover_Letter.pdf", key="pdf_cl")
                except Exception as e:
                    st.error(f"PDF Error: {e}")

                st.download_button("⬇️ TXT", data=cover_text_export, file_name="Cover_Letter.txt", key="txt_cl")
            else:
                st.warning("No cover letter available to export.")